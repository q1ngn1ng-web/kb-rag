"""DOCX 文档导入器：使用 python-docx 将 .docx 文件转为 Markdown。"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from .base import Importer, ImportResult, auto_register


@auto_register
class DocxImporter(Importer):
    """将 DOCX 文件转换为 Markdown。

    支持：
    - 标题检测（Heading 1~6 → #~######）
    - 粗体 / 斜体 格式
    - 列表项（List Bullet / List Number）
    - 表格 → Markdown 表格
    - 图片提取
    - 元数据（作者、标题、注释）
    """

    extensions = [".docx"]

    # ── 公开接口 ──────────────────────────────────────────────

    def convert(self, file_path: Path) -> ImportResult:
        try:
            doc = Document(file_path)
        except Exception as exc:
            return ImportResult(
                markdown="",
                metadata={"converter": "docx", "source_format": "docx"},
                success=False,
                error=str(exc),
            )

        metadata = self._extract_metadata(doc, file_path)
        metadata["converter"] = "docx"
        metadata["source_format"] = "docx"

        images: list[tuple[str, bytes]] = self._extract_images(doc)
        markdown = self._convert_body(doc)

        return ImportResult(markdown=markdown, metadata=metadata, images=images)

    # ── 元数据 ────────────────────────────────────────────────

    def _extract_metadata(self, doc: Document, file_path: Path) -> dict:
        meta = {}
        try:
            props = doc.core_properties
            if props.title:
                meta["title"] = props.title
            if props.author:
                meta["author"] = props.author
            if props.comments:
                meta["comments"] = props.comments
        except Exception:
            pass

        if "title" not in meta:
            meta["title"] = file_path.stem

        # 粗略估算页数（~40 段落 / 页）
        para_count = len(doc.paragraphs)
        meta["page_count_estimate"] = max(1, para_count // 40)
        return meta

    # ── 图片提取 ──────────────────────────────────────────────

    def _extract_images(self, doc: Document) -> list[tuple[str, bytes]]:
        images: list[tuple[str, bytes]] = []
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    images.append((rel.target_part.partname, rel.target_part.blob))
        except Exception:
            pass
        return images

    # ── 正文转换 ──────────────────────────────────────────────

    def _convert_body(self, doc: Document) -> str:
        """按文档顺序遍历段落和表格，生成 Markdown。"""
        lines: list[str] = []

        # 建立 id → 对象 映射以保证正确的文档顺序
        para_map = {id(p._element): p for p in doc.paragraphs}
        table_map = {id(t._element): t for t in doc.tables}

        for child in doc.element.body:
            cid = id(child)
            if cid in para_map:
                md = self._convert_paragraph(para_map[cid])
                if md is not None:
                    lines.append(md)
            elif cid in table_map:
                md = self._convert_table(table_map[cid])
                if md:
                    lines.append(md)
                    lines.append("")

        return "\n".join(lines).strip()

    # ── 段落转换 ──────────────────────────────────────────────

    def _convert_paragraph(self, para) -> str | None:
        text = para.text.strip()
        if not text:
            return None

        style_name = (para.style.name or "").lower() if para.style else ""

        # --- 标题 ---
        if style_name.startswith("heading"):
            try:
                level = int(style_name.split()[-1])
                level = min(level, 6)
                return f"{'#' * level} {text}"
            except (ValueError, IndexError):
                pass
            return f"### {text}"

        if style_name == "title":
            return f"# {text}"

        # --- 列表 ---
        prefix = ""
        is_list = style_name.startswith("list")
        if is_list:
            if "bullet" in style_name:
                prefix = "- "
            elif "number" in style_name:
                prefix = "1. "
            else:
                prefix = "- "

        formatted = self._format_runs(para)
        if is_list:
            return f"{prefix}{formatted}"
        return formatted

    # ── 行内格式 ──────────────────────────────────────────────

    def _format_runs(self, para) -> str:
        parts: list[str] = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            parts.append(text)

        return "".join(parts) if parts else para.text

    # ── 表格转换 ──────────────────────────────────────────────

    def _convert_table(self, table) -> str:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        num_cols = max(len(r) for r in rows)
        lines: list[str] = []

        # 表头（第一行）
        header = rows[0]
        while len(header) < num_cols:
            header.append("")
        lines.append("| " + " | ".join(header) + " |")

        # 分隔行
        lines.append("| " + " | ".join(["---"] * num_cols) + " |")

        # 数据行
        for row in rows[1:]:
            while len(row) < num_cols:
                row.append("")
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)
